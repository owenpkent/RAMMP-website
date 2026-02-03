"""
Convert Markdown files to Word documents (.docx).

Usage:
    python md_to_docx.py <input.md> [output.docx]
    
If output path is not specified, creates .docx in same directory as input.

Dependencies:
    pip install python-docx markdown
"""
import argparse
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def parse_markdown_to_elements(md_content: str) -> list:
    """Parse markdown content into structured elements."""
    elements = []
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Headers
        if line.startswith('#'):
            level = len(re.match(r'^#+', line).group())
            text = line.lstrip('#').strip()
            elements.append({'type': 'heading', 'level': level, 'text': text})
            i += 1
            continue
        
        # Tables
        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            elements.append({'type': 'table', 'lines': table_lines})
            continue
        
        # Bullet lists
        if re.match(r'^[\s]*[-*•]\s', line):
            list_items = []
            while i < len(lines) and re.match(r'^[\s]*[-*•]\s', lines[i]):
                indent = len(re.match(r'^[\s]*', lines[i]).group())
                text = re.sub(r'^[\s]*[-*•]\s+', '', lines[i])
                list_items.append({'indent': indent, 'text': text})
                i += 1
            elements.append({'type': 'list', 'items': list_items})
            continue
        
        # Numbered lists
        if re.match(r'^[\s]*\d+\.\s', line):
            list_items = []
            while i < len(lines) and re.match(r'^[\s]*\d+\.\s', lines[i]):
                text = re.sub(r'^[\s]*\d+\.\s+', '', lines[i])
                list_items.append({'text': text})
                i += 1
            elements.append({'type': 'numbered_list', 'items': list_items})
            continue
        
        # Regular paragraph
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith('#') and '|' not in lines[i] and not re.match(r'^[\s]*[-*•]\s', lines[i]) and not re.match(r'^[\s]*\d+\.\s', lines[i]):
            para_lines.append(lines[i])
            i += 1
        elements.append({'type': 'paragraph', 'text': ' '.join(para_lines)})
    
    return elements


def apply_inline_formatting(paragraph, text: str):
    """Apply bold and italic formatting to text within a paragraph."""
    # Pattern for **bold**, *italic*, `code`
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)'
    
    last_end = 0
    for match in re.finditer(pattern, text):
        # Add text before match
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])
        
        # Add formatted text
        if match.group(2):  # Bold
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # Italic
            run = paragraph.add_run(match.group(3))
            run.italic = True
        elif match.group(4):  # Code
            run = paragraph.add_run(match.group(4))
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        
        last_end = match.end()
    
    # Add remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def markdown_to_docx(md_path: Path, docx_path: Path = None):
    """Convert a Markdown file to a Word document."""
    if docx_path is None:
        docx_path = md_path.with_suffix('.docx')
    
    md_content = md_path.read_text(encoding='utf-8')
    elements = parse_markdown_to_elements(md_content)
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    for element in elements:
        if element['type'] == 'heading':
            level = min(element['level'], 4)  # Word has Heading 1-4
            heading = doc.add_heading(level=level)
            apply_inline_formatting(heading, element['text'])
        
        elif element['type'] == 'paragraph':
            para = doc.add_paragraph()
            apply_inline_formatting(para, element['text'])
        
        elif element['type'] == 'list':
            for item in element['items']:
                para = doc.add_paragraph(style='List Bullet')
                apply_inline_formatting(para, item['text'])
        
        elif element['type'] == 'numbered_list':
            for item in element['items']:
                para = doc.add_paragraph(style='List Number')
                apply_inline_formatting(para, item['text'])
        
        elif element['type'] == 'table':
            table_lines = element['lines']
            if len(table_lines) < 2:
                continue
            
            # Parse header row
            headers = [cell.strip() for cell in table_lines[0].split('|') if cell.strip()]
            
            # Skip separator row (index 1) and parse data rows
            data_rows = []
            for row_line in table_lines[2:]:
                cells = [cell.strip() for cell in row_line.split('|') if cell.strip()]
                if cells:
                    data_rows.append(cells)
            
            if not headers:
                continue
            
            # Create table
            table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
            table.style = 'Table Grid'
            
            # Add headers
            header_row = table.rows[0]
            for i, header in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = ''
                para = cell.paragraphs[0]
                run = para.add_run(header.replace('**', ''))
                run.bold = True
            
            # Add data rows
            for row_idx, row_data in enumerate(data_rows):
                row = table.rows[row_idx + 1]
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < len(headers):
                        cell = row.cells[col_idx]
                        cell.text = ''
                        para = cell.paragraphs[0]
                        # Handle bold in cells
                        clean_text = cell_text.replace('**', '')
                        if cell_text.startswith('**') and cell_text.endswith('**'):
                            run = para.add_run(clean_text)
                            run.bold = True
                        else:
                            para.add_run(clean_text)
            
            # Add spacing after table
            doc.add_paragraph()
    
    doc.save(str(docx_path))
    print(f"Created: {docx_path}")
    return docx_path


def main():
    parser = argparse.ArgumentParser(
        description='Convert Markdown files to Word documents',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
    python md_to_docx.py document.md
    python md_to_docx.py document.md output.docx
    python md_to_docx.py ../proposals/analysis.md ./exports/analysis.docx
        """
    )
    parser.add_argument('input', help='Input Markdown file path')
    parser.add_argument('output', nargs='?', help='Output Word document path (optional)')
    
    args = parser.parse_args()
    
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: Input file not found: {input_path}")
        return 1
    
    output_path = Path(args.output) if args.output else None
    
    try:
        markdown_to_docx(input_path, output_path)
        return 0
    except Exception as e:
        print(f"Error converting file: {e}")
        return 1


if __name__ == "__main__":
    exit(main())
